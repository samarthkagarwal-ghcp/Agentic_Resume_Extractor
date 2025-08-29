from docx import Document
from pydantic import BaseModel, Field
from crewai.tools import BaseTool
import json

class ExtractTextInput(BaseModel):
    """Input schema for DocumentTool."""
    file_path: str = Field(..., description="Path to the document file to be read. Can be absolute or relative path.")
    page_range: str = Field(
        "all",
        description="Range of pages to read. Format: 'all' for entire document, '1-5' for pages 1 to 5, '3' for just page 3."
    )

class ExtractTextTool(BaseTool):

    name: str = "Extract Text from DOCX"
    description: str = "Extracts all text from a .docx document file, including paragraphs and tables."

    def _run(self, file_path: str, page_range: str = "all") -> str:
        if not file_path.lower().endswith(".docx"):
            return json.dumps({"status": "error", "message": f"Unsupported file type for {file_path}. Only .docx is supported."})
        
        try:
            text = self.extract_text_from_docx(file_path)
            if not text.strip():
                return json.dumps({"status": "error", "message": "No text extracted from document"})
            
            return json.dumps({"status": "success", "file_path": file_path, "content": text})
        except Exception as e:
            return json.dumps({"status": "error", "message": f"Error extracting DOCX content: {e}"})

    def extract_text_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        texts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # Extract tables (row by row, cell by cell)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    texts.append(" | ".join(row_text))

        return "\n".join(texts)
